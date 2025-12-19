"""
Document Understanding Agent
Handles text extraction, OCR for images, and voice transcription cleanup
"""

import os
import re
from typing import Optional, Dict, Any, List
from pathlib import Path
import tempfile

# OCR for images
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("Warning: OCR libraries not available. Image processing will be limited.")

# Voice/audio processing
try:
    import speech_recognition as sr
    VOICE_AVAILABLE = True
except ImportError:
    VOICE_AVAILABLE = False
    print("Warning: Speech recognition not available. Voice input will not work.")

# Document processing
try:
    import PyPDF2
    import pypdf
    from docx import Document
    DOC_PROCESSING_AVAILABLE = True
except ImportError:
    DOC_PROCESSING_AVAILABLE = False

from ..config import Config


class DocumentUnderstandingAgent:
    """
    Agent responsible for understanding and cleaning input content.
    Handles:
    - Text documents (PDF, TXT, DOCX)
    - Images with OCR
    - Voice transcriptions
    """
    
    def __init__(self):
        self.config = Config()
        self.ocr_available = OCR_AVAILABLE
        self.voice_available = VOICE_AVAILABLE
        
        if self.voice_available:
            self.recognizer = sr.Recognizer()
    
    def process_input(
        self, 
        input_source: str,
        input_type: str = "text",
        file_path: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Process input content based on type
        
        Args:
            input_source: The input content (text string, file path, or audio file path)
            input_type: Type of input - "text", "image", "voice", "document"
            file_path: Optional file path for file-based inputs
            
        Returns:
            Dictionary with cleaned text and metadata
        """
        try:
            if input_type == "text":
                cleaned_text = self._clean_text(input_source)
                return {
                    "success": True,
                    "cleaned_text": cleaned_text,
                    "original_length": len(input_source),
                    "cleaned_length": len(cleaned_text),
                    "input_type": "text"
                }
            
            elif input_type == "image":
                if not self.ocr_available:
                    return {
                        "success": False,
                        "error": "OCR not available. Please install pytesseract and PIL."
                    }
                cleaned_text = self._extract_text_from_image(input_source if file_path is None else file_path)
                return {
                    "success": True,
                    "cleaned_text": cleaned_text,
                    "input_type": "image",
                    "ocr_used": True
                }
            
            elif input_type == "voice":
                if not self.voice_available:
                    return {
                        "success": False,
                        "error": "Speech recognition not available."
                    }
                cleaned_text = self._process_voice_input(input_source if file_path is None else file_path)
                return {
                    "success": True,
                    "cleaned_text": cleaned_text,
                    "input_type": "voice",
                    "transcription_used": True
                }
            
            elif input_type == "document":
                cleaned_text = self._extract_from_document(input_source if file_path is None else file_path)
                return {
                    "success": True,
                    "cleaned_text": cleaned_text,
                    "input_type": "document"
                }
            
            else:
                return {
                    "success": False,
                    "error": f"Unsupported input type: {input_type}"
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "input_type": input_type
            }
    
    def _extract_text_from_image(self, image_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            if not self.ocr_available:
                raise Exception("OCR libraries not available")
            
            image = Image.open(image_path)
            # Perform OCR
            raw_text = pytesseract.image_to_string(image)
            # Clean OCR output
            cleaned_text = self._clean_ocr_text(raw_text)
            return cleaned_text
        except pytesseract.TesseractNotFoundError:
            error_msg = (
                "Tesseract OCR is not installed. "
                "Please install Tesseract OCR:\n"
                "Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki\n"
                "After installation, add Tesseract to your PATH or set TESSDATA_PREFIX environment variable."
            )
            print(f"OCR Error: {error_msg}")
            raise Exception(error_msg)
        except Exception as e:
            error_msg = f"OCR extraction error: {str(e)}"
            print(error_msg)
            if "tesseract" in str(e).lower() or "not found" in str(e).lower():
                raise Exception(
                    "Tesseract OCR not found. Please install Tesseract OCR:\n"
                    "Windows: https://github.com/UB-Mannheim/tesseract/wiki\n"
                    "After installation, restart the application."
                )
            raise Exception(error_msg)
    
    def _process_voice_input(self, audio_path: str) -> str:
        """Process voice/audio input and return cleaned transcription"""
        try:
            with sr.AudioFile(audio_path) as source:
                audio = self.recognizer.record(source)
                # Use Google's speech recognition
                raw_text = self.recognizer.recognize_google(audio)
                # Clean transcription (remove filler words, fix spoken grammar)
                cleaned_text = self._clean_voice_text(raw_text)
                return cleaned_text
        except sr.UnknownValueError:
            print("Could not understand audio")
            return ""
        except sr.RequestError as e:
            print(f"Speech recognition error: {e}")
            return ""
    
    def _extract_from_document(self, file_path: str) -> str:
        """Extract text from document files (PDF, DOCX, TXT)"""
        file_ext = Path(file_path).suffix.lower()
        text = ""
        
        try:
            if file_ext == '.pdf':
                text = self._extract_from_pdf(file_path)
            elif file_ext == '.txt':
                text = self._extract_from_txt(file_path)
            elif file_ext == '.docx':
                text = self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported document format: {file_ext}")
            
            return self._clean_text(text)
        except Exception as e:
            print(f"Document extraction error: {e}")
            return ""
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        
        # Try PyPDF2 first
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"PyPDF2 extraction failed: {e}")
        
        # Fallback to pypdf
        if len(text.strip()) < 100:
            try:
                text = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                print(f"pypdf extraction failed: {e}")
        
        return text
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                print(f"Text file reading error: {e}")
                return ""
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text += row_text + "\n"
            
            return text
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove common PDF/document artifacts
        artifacts = [
            '\x0c',  # Form feed
            '\x0b',  # Vertical tab
            'Â',     # Encoding artifact
            'â€™',   # Quote encoding
            'â€œ',   # Quote encoding
            'â€',   # Quote encoding
        ]
        
        for artifact in artifacts:
            text = text.replace(artifact, '')
        
        # Fix common quote issues
        text = text.replace('â€™', "'").replace('â€œ', '"').replace('â€', '"')
        
        # Remove excessive newlines but preserve paragraph breaks
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        text = '\n\n'.join(lines)
        
        return text.strip()
    
    def _clean_ocr_text(self, text: str) -> str:
        """Clean OCR-extracted text, correcting common OCR errors"""
        if not text:
            return ""
        
        # Basic cleaning first
        text = self._clean_text(text)
        
        # Common OCR error corrections (context-aware)
        ocr_corrections = {
            r'\b0\b': 'o',  # Zero mistaken for 'o' in words (context-dependent)
            r'rn\b': 'm',   # 'rn' mistaken for 'm' at word end
            r'vv': 'w',     # 'vv' mistaken for 'w'
            r'ii\b': 'n',   # 'ii' mistaken for 'n' at word end
        }
        
        # Apply corrections (be careful with these)
        # This is a simplified version - real OCR correction would use context
        
        # Fix spacing issues
        text = re.sub(r'\s+', ' ', text)
        
        # Fix line breaks in the middle of sentences
        text = re.sub(r'([a-z])\n([a-z])', r'\1 \2', text)
        
        return text.strip()
    
    def _clean_voice_text(self, text: str) -> str:
        """Clean voice transcription, removing filler words and fixing spoken grammar"""
        if not text:
            return ""
        
        # Common filler words to remove
        filler_words = [
            'um', 'uh', 'er', 'ah', 'like', 'you know', 'I mean',
            'so', 'well', 'actually', 'basically', 'literally'
        ]
        
        # Remove filler words (with word boundaries)
        for filler in filler_words:
            # Case-insensitive removal with word boundaries
            pattern = r'\b' + re.escape(filler) + r'\b'
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)
        
        # Fix common spoken grammar issues
        # "gonna" -> "going to"
        text = re.sub(r'\bgonna\b', 'going to', text, flags=re.IGNORECASE)
        text = re.sub(r'\bwanna\b', 'want to', text, flags=re.IGNORECASE)
        text = re.sub(r'\bgotta\b', 'got to', text, flags=re.IGNORECASE)
        
        # Remove excessive spaces
        text = re.sub(r'\s+', ' ', text)
        
        # Capitalize sentences
        sentences = re.split(r'([.!?]\s+)', text)
        cleaned_sentences = []
        for i, sentence in enumerate(sentences):
            if i == 0 or (i > 0 and sentences[i-1].strip().endswith(('.', '!', '?'))):
                sentence = sentence.capitalize()
            cleaned_sentences.append(sentence)
        text = ''.join(cleaned_sentences)
        
        # Fix double spaces
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    def extract_key_insights(self, text: str) -> Dict[str, Any]:
        """
        Extract key insights from cleaned text
        
        Args:
            text: Cleaned text content
            
        Returns:
            Dictionary with key insights (topics, themes, etc.)
        """
        # Simple keyword and topic extraction
        words = text.lower().split()
        
        # Get word frequencies (excluding common stop words)
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 
                     'to', 'for', 'of', 'with', 'by', 'from', 'as', 'is', 'was',
                     'are', 'were', 'be', 'been', 'have', 'has', 'had', 'do',
                     'does', 'did', 'will', 'would', 'could', 'should', 'may',
                     'might', 'this', 'that', 'these', 'those', 'i', 'you', 'he',
                     'she', 'it', 'we', 'they'}
        
        word_freq = {}
        for word in words:
            # Remove punctuation
            word = re.sub(r'[^\w]', '', word)
            if word and word not in stop_words and len(word) > 2:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Get top keywords
        top_keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
        
        # Estimate topic count (simple heuristic)
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        estimated_topics = min(max(3, len(paragraphs) // 3), 20)
        
        return {
            "text_length": len(text),
            "paragraph_count": len(paragraphs),
            "estimated_topics": estimated_topics,
            "top_keywords": [word for word, freq in top_keywords],
            "word_count": len(words)
        }
