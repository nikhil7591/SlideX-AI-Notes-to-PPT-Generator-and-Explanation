"""
Document Processing Module
Handles text extraction from PDF, TXT, and DOCX files
"""

import os
import tempfile
from typing import Optional
from pathlib import Path

# PDF Processing
import PyPDF2
import pypdf

# Text Processing
from docx import Document

class DocumentProcessor:
    """Handles document text extraction and preprocessing"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.txt', '.docx']
    
    def extract_text(self, file_path: str, file_extension: str) -> Optional[str]:
        """
        Extract text from uploaded document
        
        Args:
            file_path: Path to the uploaded file
            file_extension: File extension (.pdf, .txt, .docx)
            
        Returns:
            Extracted text as string or None if extraction fails
        """
        try:
            file_extension = file_extension.lower()
            
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            print(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using multiple methods"""
        text = ""
        
        # Try PyPDF2 first
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"PyPDF2 failed: {e}")
        
        # If PyPDF2 didn't work well, try pypdf
        if len(text.strip()) < 100:  # If very little text extracted
            try:
                text = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                print(f"pypdf failed: {e}")
        
        return self._clean_text(text)
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return self._clean_text(text)
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                return self._clean_text(text)
            except Exception as e:
                print(f"Error reading text file: {e}")
                return ""
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text += row_text + "\n"
            
            return self._clean_text(text)
        except Exception as e:
            print(f"Error extracting from DOCX: {e}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean and preprocess extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove common PDF artifacts
        artifacts = [
            '\x0c',  # Form feed
            '\x0b',  # Vertical tab
            'Â',     # Common encoding artifact
            'â€™',   # Quote encoding artifact
            'â€œ',   # Quote encoding artifact
            'â€',   # Quote encoding artifact
        ]
        
        for artifact in artifacts:
            text = text.replace(artifact, '')
        
        # Fix common quote issues
        text = text.replace('â€™', "'").replace('â€œ', '"').replace('â€', '"')
        
        # Remove excessive newlines
        text = '\n'.join([line.strip() for line in text.split('\n') if line.strip()])
        
        return text.strip()
    
    def chunk_text(self, text: str, chunk_size: int = 4000, overlap: int = 200) -> list:
        """
        Split text into chunks for processing
        
        Args:
            text: Text to chunk
            chunk_size: Maximum size of each chunk
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        if not text:
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            if end >= len(text):
                chunks.append(text[start:])
                break
            
            # Try to break at a sentence or paragraph
            chunk = text[start:end]
            
            # Look for sentence endings
            sentence_endings = ['. ', '! ', '? ', '\n']
            best_break = -1
            
            for ending in sentence_endings:
                last_occurrence = chunk.rfind(ending)
                if last_occurrence > best_break:
                    best_break = last_occurrence + len(ending)
            
            if best_break > 0:
                chunks.append(text[start:start + best_break])
                start = start + best_break - overlap
            else:
                # If no good break point, use the chunk as is
                chunks.append(chunk)
                start = end - overlap
        
        return [chunk.strip() for chunk in chunks if chunk.strip()]
    
    def get_document_info(self, file_path: str) -> dict:
        """Get basic information about the document"""
        try:
            file_size = os.path.getsize(file_path)
            file_ext = Path(file_path).suffix.lower()
            
            info = {
                'file_path': file_path,
                'file_size': file_size,
                'file_extension': file_ext,
                'file_size_mb': round(file_size / (1024 * 1024), 2)
            }
            
            # Extract text to get word count
            text = self.extract_text(file_path, file_ext)
            if text:
                info['word_count'] = len(text.split())
                info['char_count'] = len(text)
                info['text_preview'] = text[:500] + "..." if len(text) > 500 else text
            else:
                info['word_count'] = 0
                info['char_count'] = 0
                info['text_preview'] = ""
            
            return info
            
        except Exception as e:
            print(f"Error getting document info: {e}")
            return {
                'file_path': file_path,
                'error': str(e)
            }
